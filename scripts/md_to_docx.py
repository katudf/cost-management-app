# -*- coding: utf-8 -*-
"""docs/design.md を design.docx に変換するワンオフスクリプト。"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "docs/design.md"
DST = "docs/design.docx"


def set_japanese_font(run, size=10.5, bold=False, name="Yu Gothic"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_japanese_font(run, size={1: 18, 2: 15, 3: 13, 4: 11.5}.get(level, 11), bold=True)
    return h


def add_paragraph(doc, text, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    for part, is_code in split_inline_code(text):
        run = p.add_run(part)
        set_japanese_font(run, bold=False)
        run.italic = italic
        if is_code:
            run.font.name = "Consolas"
            run.font.highlight_color = 7  # gray-ish, index for WD_COLOR_INDEX.GRAY_25 not directly; skip if fails
    return p


def split_inline_code(text):
    """Split text into (segment, is_code) tuples on `code` spans."""
    parts = []
    pattern = re.compile(r'`([^`]+)`')
    idx = 0
    for m in pattern.finditer(text):
        if m.start() > idx:
            parts.append((strip_md_inline(text[idx:m.start()]), False))
        parts.append((m.group(1), True))
        idx = m.end()
    if idx < len(text):
        parts.append((strip_md_inline(text[idx:]), False))
    if not parts:
        parts = [(strip_md_inline(text), False)]
    return parts


def strip_md_inline(text):
    # 太字・リンクのMarkdown記法をプレーンテキスト化
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', text)
    return text


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    text = "\n".join(lines) if lines else " "
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), "Consolas")
    return p


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(strip_md_inline(h))
        set_japanese_font(run, bold=True, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cells):
                break
            cells[i].text = ""
            for part, is_code in split_inline_code(val):
                run = cells[i].paragraphs[0].add_run(part)
                set_japanese_font(run, size=9.5)
                if is_code:
                    run.font.name = "Consolas"
    doc.add_paragraph()
    return table


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def is_separator_row(cells):
    return all(re.fullmatch(r':?-+:?', c) for c in cells if c != '')


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Yu Gothic"

    i = 0
    n = len(lines)
    in_code = False
    code_lines = []
    code_fence_lang = None

    while i < n:
        line = lines[i]

        # コードブロック（mermaid含む）
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
                code_fence_lang = line.strip()[3:].strip()
            else:
                in_code = False
                add_code_block(doc, code_lines)
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # 水平線
        if re.fullmatch(r'-{3,}', stripped):
            i += 1
            continue

        # 見出し
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            text = strip_md_inline(m.group(2))
            add_heading(doc, text, min(level, 4))
            i += 1
            continue

        # 引用（> ...）
        if stripped.startswith('>'):
            quote_lines = []
            while i < n and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].strip()))
                i += 1
            p = add_paragraph(doc, " ".join(quote_lines), italic=True)
            p.paragraph_format.left_indent = Cm(0.75)
            continue

        # テーブル
        if stripped.startswith('|') and i + 1 < n and is_separator_row(parse_table_row(lines[i + 1])):
            header = parse_table_row(lines[i])
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # 番号付きリスト
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            add_paragraph(doc, strip_md_inline_keep_links(m.group(2)), style="List Number")
            i += 1
            continue

        # 箇条書き（インデント考慮）
        m = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if m:
            indent = len(m.group(1))
            text = m.group(2)
            style_name = "List Bullet2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style_name)
            for part, is_code in split_inline_code(text):
                run = p.add_run(part)
                set_japanese_font(run)
                if is_code:
                    run.font.name = "Consolas"
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 通常段落
        add_paragraph(doc, stripped)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


def strip_md_inline_keep_links(text):
    return strip_md_inline(text)


if __name__ == "__main__":
    main()
